"""
Medical Data Extractor - استخراج البيانات الطبية من أي ملف
Supports: Images, PDF, Excel, Word, Text
Output: Same DataFrame format as the original genetic dataset
"""

import pytesseract
import pdfplumber
import pandas as pd
import numpy as np
import re
import json
import os
from PIL import Image
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

# تثبيت المكتبات المطلوبة
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available. Install with: pip install python-docx")

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("⚠️ google-generativeai not available. Install with: pip install google-generativeai")

# ==================== إعداد Gemini ====================
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyBnUf97bdEt_WlLdC79LNX1lqRak1d5s_Y")

if GEMINI_AVAILABLE and GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    gemini_model = genai.GenerativeModel("gemini-1.5-flash")
    print("✅ Gemini AI is ready!")
else:
    gemini_model = None
    print("⚠️ Gemini AI not configured")

# ==================== استخراج النص من الملفات ====================

def extract_text_from_file(file_path):
    """
    استخراج النص من أي نوع ملف
    
    Supported formats:
    - Images: .jpg, .png, .jpeg
    - PDF: .pdf
    - Excel: .xlsx, .xls
    - Word: .docx
    - Text: .txt
    """
    ext = Path(file_path).suffix.lower()
    text = ""
    
    try:
        if ext in ['.jpg', '.png', '.jpeg']:
            # استخراج من الصور باستخدام OCR
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang='eng+ara')
            print(f"   📸 OCR from image: {len(text)} chars")
        
        elif ext == '.pdf':
            # استخراج من PDF
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            print(f"   📄 PDF: {len(pdf.pages)} pages, {len(text)} chars")
        
        elif ext in ['.xlsx', '.xls']:
            # قراءة Excel مباشرة
            df = pd.read_excel(file_path)
            text = df.to_string()
            print(f"   📊 Excel: {df.shape[0]} rows, {df.shape[1]} cols")
        
        elif ext == '.docx' and DOCX_AVAILABLE:
            # قراءة Word
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            print(f"   📝 Word: {len(doc.paragraphs)} paragraphs")
        
        elif ext == '.txt':
            # قراءة ملف نصي
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            print(f"   📃 Text file: {len(text)} chars")
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    except Exception as e:
        print(f"   ❌ Error: {e}")
        return ""
    
    return text

# ==================== استخراج البيانات باستخدام Regex ====================

def extract_values_with_regex(text):
    """استخراج القيم الطبية باستخدام Regular Expressions"""
    data = {}
    
    patterns = {
        'age': r'(?:age|عمر|Age)[\s:]*(\d+)',
        'glucose': r'(?:glucose|سكر|Glucose|blood sugar)[\s:]*(\d+(?:\.\d+)?)',
        'systolic_bp': r'(?:systolic|الضغط الانقباضي|Systolic)[\s:]*(\d+(?:\.\d+)?)',
        'diastolic_bp': r'(?:diastolic|الضغط الانبساطي|Diastolic)[\s:]*(\d+(?:\.\d+)?)',
        'ldl': r'(?:ldl|LDL)[\s:]*(\d+(?:\.\d+)?)',
        'genetic_risk_score': r'(?:genetic risk|الخطر الوراثي)[\s:]*(\d+(?:\.\d+)?)',
        'mutation_load': r'(?:mutation load|حمولة الطفرة)[\s:]*(\d+(?:\.\d+)?)',
        'penetrance': r'(?:penetrance|الاختراق)[\s:]*(\d+(?:\.\d+)?)'
    }
    
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                val = match.group(1)
                data[key] = float(val) if '.' in val else int(val)
            except:
                data[key] = match.group(1)
    
    # Gender
    if re.search(r'male|ذكر|Male', text, re.IGNORECASE):
        data['gender'] = 'Male'
    elif re.search(r'female|انثى|Female', text, re.IGNORECASE):
        data['gender'] = 'Female'
    
    # Genetic disease
    disease_match = re.search(r'(?:genetic disease|مرض وراثي|Diagnosis)[\s:]*([A-Za-z\s]+)', text, re.IGNORECASE)
    if disease_match:
        data['genetic_disease'] = disease_match.group(1).strip()
    
    # Inheritance pattern
    inh_match = re.search(r'(?:inheritance|وراثة|Inheritance)[\s:]*([A-Za-z\s]+)', text, re.IGNORECASE)
    if inh_match:
        data['inheritance_pattern'] = inh_match.group(1).strip()
    
    return data

# ==================== استخراج البيانات باستخدام Gemini AI ====================

def extract_values_with_gemini(text):
    """استخراج البيانات الطبية باستخدام Gemini AI"""
    if gemini_model is None:
        return {}
    
    prompt = f"""
    Extract medical values from this text. Return ONLY valid JSON.
    
    Extract:
    - age (int)
    - glucose (float, mg/dL)
    - systolic_bp (float, mmHg)
    - diastolic_bp (float, mmHg)
    - ldl (float, mg/dL)
    - genetic_risk_score (float, 0-1)
    - genetic_disease (string)
    - inheritance_pattern (string)
    - gender (string: "Male" or "Female")
    
    Text: {text[:3000]}
    
    Output example: {{"age": 45, "glucose": 110, "systolic_bp": 120, "diastolic_bp": 80, "ldl": 130, "genetic_risk_score": 0.5, "genetic_disease": "Diabetes Type 2", "inheritance_pattern": "Autosomal Dominant", "gender": "Male"}}
    """
    
    try:
        response = gemini_model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Clean JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        
        return json.loads(response_text)
    except Exception as e:
        print(f"   ⚠️ Gemini error: {e}")
        return {}

# ==================== تحويل النص إلى DataFrame ====================

def text_to_dataset(text, use_gemini=True, fill_missing=True):
    """
    تحويل النص إلى DataFrame بنفس تنسيق الـ dataset الأصلي
    """
    
    # استخراج البيانات
    data = extract_values_with_regex(text)
    
    if use_gemini and gemini_model:
        gemini_data = extract_values_with_gemini(text)
        data.update(gemini_data)
    
    # الأعمدة المطلوبة (نفس الـ dataset الأصلي)
    columns = [
        'person_id', 'family_id', 'family_name', 'name', 'relation',
        'generation', 'age', 'gender', 'parent_id',
        'genetic_risk_score', 'genetic_disease', 'affected_gene',
        'inheritance_pattern', 'carrier_status', 'mutation_load',
        'penetrance', 'age_of_onset', 'family_mutations',
        'glucose', 'systolic_bp', 'diastolic_bp', 'ldl',
        'health_condition', 'disease_manifested'
    ]
    
    if not fill_missing:
        available = [c for c in columns if c in data]
        return pd.DataFrame([{c: data.get(c) for c in available}]) if available else pd.DataFrame()
    
    # قيم افتراضية
    default_values = {
        'person_id': f"P{np.random.randint(100000, 999999)}",
        'family_id': f"F{np.random.randint(100000, 999999)}",
        'family_name': 'Unknown',
        'name': 'Unknown',
        'relation': 'Unknown',
        'generation': 1,
        'age': 40,
        'gender': 'Unknown',
        'parent_id': None,
        'genetic_risk_score': 0.3,
        'genetic_disease': 'None',
        'affected_gene': None,
        'inheritance_pattern': 'None',
        'carrier_status': 0,
        'mutation_load': 0.0,
        'penetrance': 0.0,
        'age_of_onset': None,
        'family_mutations': 'None',
        'glucose': 0.0,
        'systolic_bp': 0.0,
        'diastolic_bp': 0.0,
        'ldl': 0.0,
        'health_condition': 'None',
        'disease_manifested': 0
    }
    
    row = {}
    for col in columns:
        if col in data and data[col] is not None:
            row[col] = data[col]
        else:
            row[col] = default_values[col]
    
    return pd.DataFrame([row])

# ==================== الوظيفة الرئيسية ====================

def file_to_dataset(file_path, use_gemini=True, fill_missing=True, verbose=True):
    """
    الوظيفة الرئيسية: أي ملف → DataFrame بنفس تنسيق التدريب
    
    Parameters:
    -----------
    file_path: str - مسار الملف (صورة، PDF، Excel، Word، نص)
    use_gemini: bool - استخدام Gemini AI للتحسين
    fill_missing: bool - ملء القيم المفقودة
    verbose: bool - طباعة التفاصيل
    
    Returns:
    --------
    pd.DataFrame - البيانات بالتنسيق المطلوب
    """
    
    if verbose:
        print(f"\n{'='*50}")
        print(f"📂 Processing: {file_path}")
        print('='*50)
    
    # استخراج النص من الملف
    text = extract_text_from_file(file_path)
    
    if not text:
        print("❌ Failed to extract text")
        return pd.DataFrame()
    
    if verbose:
        print(f"✅ Extracted {len(text)} characters")
        print("🔄 Converting to dataset...")
    
    # تحويل النص إلى DataFrame
    df = text_to_dataset(text, use_gemini, fill_missing)
    
    if verbose and not df.empty:
        print("\n📊 Result:")
        display_cols = ['person_id', 'age', 'gender', 'genetic_disease', 
                        'glucose', 'systolic_bp', 'ldl']
        display_cols = [c for c in display_cols if c in df.columns]
        print(df[display_cols].to_string(index=False))
        print(f"\n✅ Shape: {df.shape}")
        print(f"✅ Columns: {list(df.columns)}")
    
    return df

# ==================== دوال مساعدة إضافية ====================

def save_dataset(df, output_path):
    """حفظ DataFrame كملف CSV"""
    df.to_csv(output_path, index=False)
    print(f"✅ Saved to: {output_path}")
    return output_path

def batch_process(folder_path, use_gemini=True, output_folder=None):
    """معالجة جميع الملفات في مجلد واحد"""
    folder = Path(folder_path)
    if output_folder is None:
        output_folder = folder / "extracted_data"
    else:
        output_folder = Path(output_folder)
    
    output_folder.mkdir(exist_ok=True)
    
    all_dfs = []
    extensions = ['*.jpg', '*.png', '*.jpeg', '*.pdf', '*.xlsx', '*.xls', '*.docx', '*.txt']
    
    for ext in extensions:
        for file_path in folder.glob(ext):
            df = file_to_dataset(str(file_path), use_gemini=use_gemini, verbose=True)
            if not df.empty:
                all_dfs.append(df)
                save_dataset(df, output_folder / f"{file_path.stem}_extracted.csv")
    
    if all_dfs:
        combined = pd.concat(all_dfs, ignore_index=True)
        save_dataset(combined, output_folder / "all_extracted_data.csv")
        return combined
    
    print("No files processed")
    return pd.DataFrame()

# ==================== مثال الاستخدام ====================

if __name__ == "__main__":
    print("=" * 60)
    print("🧬 Medical Data Extractor")
    print("=" * 60)
    print("Converts any file to genetic dataset format")
    print("=" * 60)
    
    # إنشاء ملف اختبار
    test_text = """
    Patient Name: Ahmed Mohamed
    Age: 58
    Blood Results:
    - Glucose: 165 mg/dL
    - Blood Pressure: 145/90
    - LDL: 170 mg/dL
    Genetic Risk Score: 0.7
    Genetic Disease: Diabetes Type 2
    Inheritance Pattern: Autosomal Dominant
    """
    
    with open("sample_medical.txt", "w", encoding='utf-8') as f:
        f.write(test_text)
    
    print("\n✅ Created test file: sample_medical.txt")
    
    # معالجة الملف
    df = file_to_dataset("sample_medical.txt", use_gemini=False, verbose=True)
    
    # حفظ النتيجة
    if not df.empty:
        save_dataset(df, "extracted_medical_data.csv")
        print("\n✅ Done! Check 'extracted_medical_data.csv'")
    
    print("\n" + "=" * 60)
    print("🔄 Usage examples:")
    print("  df = file_to_dataset('your_file.pdf')")
    print("  df = file_to_dataset('your_image.jpg')")
    print("  df = file_to_dataset('your_data.xlsx')")
    print("  combined_df = batch_process('folder_path/')")
    print("=" * 60)